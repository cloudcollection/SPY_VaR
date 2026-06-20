from __future__ import annotations

import re
from pathlib import Path

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from utils import PROJECT_ROOT


SOURCE_MD = PROJECT_ROOT / "report" / "section_3_historical_simulation_final.md"
OUTPUT_PDF = Path(__file__).resolve().parents[1] / "pdf" / "section_3_historical_simulation_final_english.pdf"
OUTPUT_DOCX = Path(__file__).resolve().parents[1] / "outputs" / "section_3_historical_simulation_final_english.docx"
FORMULA_DIR = PROJECT_ROOT / "outputs" / "formula_images"

TIMES = "Times New Roman"
TIMES_REGULAR = "Times-Roman"
TIMES_BOLD = "Times-Bold"
PDF_BODY_FONT = "Times-Roman-Custom"
PDF_BOLD_FONT = "Times-Bold-Custom"
DOC_EAST_ASIA_FONT = TIMES
DOCUMENT_TITLE = "Section 3 Historical Simulation Methods"


def clean_text(text: str) -> str:
    return (
        text.replace("\\alpha", "alpha")
        .replace("\\lambda", "lambda")
        .replace("\\mathrm", "")
        .replace("\\widehat", "hat")
        .replace("\\infty", "infinity")
        .replace("\\leq", "<=")
        .replace("\\geq", ">=")
        .replace("\\{", "{")
        .replace("\\}", "}")
        .replace("\\", "")
    )


def clean_inline_pdf(text: str) -> str:
    text = clean_text(text)
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"`([^`]+)`", r"<font name='Times-Roman-Custom'>\1</font>", text)
    return text


def strip_markdown(text: str) -> str:
    text = clean_text(text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def parse_markdown_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    table_lines = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows = []
    for line in table_lines:
        cells = [strip_markdown(cell.strip()) for cell in line.strip("|").split("|")]
        if all(set(cell) <= {"-", ":"} for cell in cells):
            continue
        rows.append(cells)
    return rows, i


def split_markdown(md_text: str):
    lines = md_text.splitlines()
    blocks = []
    i = 0
    in_code = False
    paragraph: list[str] = []

    def flush_paragraph() -> None:
        if paragraph:
            text = " ".join(p.strip() for p in paragraph if p.strip())
            if text:
                blocks.append(("paragraph", text))
            paragraph.clear()

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_paragraph()
            in_code = not in_code
            i += 1
            continue

        if in_code:
            if stripped:
                blocks.append(("formula", stripped))
            i += 1
            continue

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            flush_paragraph()
            caption, image_path = image_match.groups()
            blocks.append(("image", (caption, image_path)))
            i += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            rows, next_i = parse_markdown_table(lines, i)
            if rows:
                blocks.append(("table", rows))
            i = next_i
            continue

        if stripped.startswith("## "):
            flush_paragraph()
            blocks.append(("heading2", stripped[3:]))
        elif stripped.startswith("### "):
            flush_paragraph()
            blocks.append(("heading3", stripped[4:]))
        elif stripped.startswith("#### "):
            flush_paragraph()
            blocks.append(("heading4", stripped[5:]))
        elif stripped.startswith("- "):
            flush_paragraph()
            blocks.append(("bullet", stripped[2:]))
        elif stripped.startswith("$$"):
            flush_paragraph()
            formula_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("$$"):
                if lines[i].strip():
                    formula_lines.append(lines[i].strip())
                i += 1
            if formula_lines:
                blocks.append(("formula", " ".join(formula_lines)))
        else:
            paragraph.append(stripped)
        i += 1

    flush_paragraph()
    return blocks


def normalize_formula(formula: str) -> str:
    formula = formula.strip().rstrip(".")
    formula = formula.replace(r"\qquad", r"\quad")
    formula = re.sub(r"\\tag\{([^}]+)\}", r"\\quad (\g<1>)", formula)
    formula = re.sub(r"\s+", " ", formula)
    return formula


def render_formula_png(formula: str, index: int) -> Path:
    FORMULA_DIR.mkdir(parents=True, exist_ok=True)
    output_path = FORMULA_DIR / f"formula_{index:02d}.png"
    latex = normalize_formula(formula)
    plt.rcParams.update(
        {
            "mathtext.fontset": "custom",
            "mathtext.rm": "Times New Roman",
            "mathtext.it": "Times New Roman:italic",
            "mathtext.bf": "Times New Roman:bold",
            "mathtext.default": "it",
            "font.family": "Times New Roman",
        }
    )
    fig = plt.figure(figsize=(7.2, 0.72), dpi=320)
    fig.patch.set_alpha(0)
    try:
        fig.text(0.01, 0.5, f"${latex}$", fontsize=11, fontfamily=TIMES, fontstyle="italic", va="center")
        fig.savefig(output_path, bbox_inches="tight", pad_inches=0.08, transparent=True)
    except Exception:
        plt.close(fig)
        # Fall back to readable monospace text if matplotlib mathtext cannot parse a formula.
        fig = plt.figure(figsize=(7.2, 0.72), dpi=320)
        fig.patch.set_alpha(0)
        fig.text(0.01, 0.5, formula, fontsize=11, fontfamily=TIMES, fontstyle="italic", va="center")
        fig.savefig(output_path, bbox_inches="tight", pad_inches=0.08, transparent=True)
    finally:
        plt.close(fig)
    return output_path


def formula_image_size(path: Path, max_width_points: float) -> tuple[float, float]:
    with PILImage.open(path) as image:
        width_px, height_px = image.size
    width = min(max_width_points, width_px * 72 / 320)
    height = width * height_px / width_px
    return width, height


def register_pdf_fonts() -> None:
    if "Times-Roman-Custom" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("Times-Roman-Custom", TIMES_REGULAR))
    if "Times-Bold-Custom" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("Times-Bold-Custom", TIMES_BOLD))
    if "STSong-Light" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))


def build_pdf(blocks) -> None:
    register_pdf_fonts()
    OUTPUT_PDF.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="TNRBody",
            parent=styles["BodyText"],
            fontName=PDF_BODY_FONT,
            fontSize=10,
            leading=13,
            alignment=TA_LEFT,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TNRSmall",
            parent=styles["BodyText"],
            fontName=PDF_BODY_FONT,
            fontSize=8,
            leading=10,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TNRH2",
            parent=styles["Heading2"],
            fontName=PDF_BOLD_FONT,
            fontSize=15,
            leading=18,
        )
    )
    styles.add(
        ParagraphStyle(
            name="TNRH3",
            parent=styles["Heading3"],
            fontName=PDF_BOLD_FONT,
            fontSize=12,
            leading=15,
        )
    )

    story = []
    formula_index = 0
    for kind, payload in blocks:
        if kind == "heading2":
            story.append(Paragraph(clean_inline_pdf(payload), styles["TNRH2"]))
            story.append(Spacer(1, 0.14 * cm))
        elif kind in {"heading3", "heading4"}:
            story.append(Paragraph(clean_inline_pdf(payload), styles["TNRH3"]))
            story.append(Spacer(1, 0.12 * cm))
        elif kind == "paragraph":
            story.append(Paragraph(clean_inline_pdf(payload), styles["TNRBody"]))
            story.append(Spacer(1, 0.1 * cm))
        elif kind == "bullet":
            story.append(Paragraph(clean_inline_pdf("• " + payload), styles["TNRBody"]))
        elif kind == "formula":
            formula_index += 1
            formula_path = render_formula_png(payload, formula_index)
            width, height = formula_image_size(formula_path, max_width_points=16.2 * cm)
            story.append(RLImage(str(formula_path), width=width, height=height))
            story.append(Spacer(1, 0.08 * cm))
        elif kind == "table":
            para_rows = [[Paragraph(clean_inline_pdf(cell), styles["TNRSmall"]) for cell in row] for row in payload]
            table = Table(para_rows, repeatRows=1)
            # Three-line table: top rule, header bottom rule, bottom rule.
            table.setStyle(
                TableStyle(
                    [
                        ("LINEABOVE", (0, 0), (-1, 0), 1.0, colors.black),
                        ("LINEBELOW", (0, 0), (-1, 0), 0.8, colors.black),
                        ("LINEBELOW", (0, -1), (-1, -1), 1.0, colors.black),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("FONTNAME", (0, 0), (-1, 0), PDF_BOLD_FONT),
                        ("LEFTPADDING", (0, 0), (-1, -1), 3),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 0.18 * cm))
        elif kind == "image":
            caption, image_path = payload
            resolved = (SOURCE_MD.parent / image_path).resolve()
            if resolved.exists():
                width, height = formula_image_size(resolved, max_width_points=16.2 * cm)
                story.append(RLImage(str(resolved), width=width, height=height))
                story.append(Paragraph(clean_inline_pdf(caption), styles["TNRSmall"]))
                story.append(Spacer(1, 0.18 * cm))

    doc = SimpleDocTemplate(
        str(OUTPUT_PDF),
        pagesize=A4,
        rightMargin=1.35 * cm,
        leftMargin=1.35 * cm,
        topMargin=1.25 * cm,
        bottomMargin=1.25 * cm,
        title=DOCUMENT_TITLE,
    )
    doc.build(story)


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = TIMES
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_EAST_ASIA_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_text(paragraph, text: str, size: float = 11, bold: bool = False) -> None:
    run = paragraph.add_run(strip_markdown(text))
    set_run_font(run, size=size, bold=bold)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(strip_markdown(text))
    set_run_font(run, size=8.5, bold=bold)


def set_cell_border(cell, **kwargs) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        if edge_data:
            for key, value in edge_data.items():
                element.set(qn(f"w:{key}"), str(value))
        else:
            element.set(qn("w:val"), "nil")


def apply_three_line_table(table) -> None:
    row_count = len(table.rows)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            borders = {}
            if r_idx == 0:
                borders["top"] = {"val": "single", "sz": "12", "color": "000000"}
                borders["bottom"] = {"val": "single", "sz": "8", "color": "000000"}
            if r_idx == row_count - 1:
                borders["bottom"] = {"val": "single", "sz": "12", "color": "000000"}
            set_cell_border(cell, **borders)


def build_docx(blocks) -> None:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = TIMES
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_EAST_ASIA_FONT)
    styles["Normal"].font.size = Pt(11)

    formula_index = 0
    for kind, payload in blocks:
        if kind == "heading2":
            p = doc.add_paragraph()
            set_paragraph_text(p, payload, size=15, bold=True)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        elif kind in {"heading3", "heading4"}:
            p = doc.add_paragraph()
            set_paragraph_text(p, payload, size=12, bold=True)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
        elif kind == "paragraph":
            p = doc.add_paragraph()
            set_paragraph_text(p, payload, size=11)
            p.paragraph_format.first_line_indent = Inches(0.25)
            p.paragraph_format.line_spacing = 1.08
            p.paragraph_format.space_after = Pt(3)
        elif kind == "bullet":
            p = doc.add_paragraph(style=None)
            set_paragraph_text(p, "• " + payload, size=11)
            p.paragraph_format.left_indent = Inches(0.2)
        elif kind == "formula":
            formula_index += 1
            formula_path = render_formula_png(payload, formula_index)
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(str(formula_path), width=Inches(5.9))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(3)
        elif kind == "table":
            rows = payload
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.autofit = True
            for r_idx, row in enumerate(rows):
                for c_idx, value in enumerate(row):
                    set_cell_text(table.cell(r_idx, c_idx), value, bold=(r_idx == 0))
            apply_three_line_table(table)
            doc.add_paragraph()
        elif kind == "image":
            caption, image_path = payload
            resolved = (SOURCE_MD.parent / image_path).resolve()
            if resolved.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(resolved), width=Inches(6.0))
                cp = doc.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_text(cp, caption, size=9)

    doc.save(OUTPUT_DOCX)


def export_document(source_md: Path, output_pdf: Path, output_docx: Path | None, title: str, chinese: bool = False) -> None:
    global SOURCE_MD, OUTPUT_PDF, OUTPUT_DOCX, PDF_BODY_FONT, PDF_BOLD_FONT, DOC_EAST_ASIA_FONT, DOCUMENT_TITLE
    SOURCE_MD = source_md
    OUTPUT_PDF = output_pdf
    OUTPUT_DOCX = output_docx
    DOCUMENT_TITLE = title
    if chinese:
        PDF_BODY_FONT = "STSong-Light"
        PDF_BOLD_FONT = "STSong-Light"
        DOC_EAST_ASIA_FONT = "SimSun"
    else:
        PDF_BODY_FONT = "Times-Roman-Custom"
        PDF_BOLD_FONT = "Times-Bold-Custom"
        DOC_EAST_ASIA_FONT = TIMES

    md_text = SOURCE_MD.read_text(encoding="utf-8")
    blocks = split_markdown(md_text)
    build_pdf(blocks)
    print(f"Exported PDF to {OUTPUT_PDF}")
    if output_docx is not None:
        build_docx(blocks)
        print(f"Exported DOCX to {OUTPUT_DOCX}")


def main() -> None:
    export_document(
        PROJECT_ROOT / "report" / "section_3_historical_simulation_final.md",
        Path(__file__).resolve().parents[1] / "pdf" / "section_3_historical_simulation_research_level_english.pdf",
        None,
        "Section 3 Historical Simulation Methods",
        chinese=False,
    )
    chinese_md = PROJECT_ROOT / "report" / "section_3_historical_simulation_final_zh.md"
    if chinese_md.exists():
        export_document(
            chinese_md,
            Path(__file__).resolve().parents[1] / "pdf" / "section_3_historical_simulation_research_level_chinese.pdf",
            None,
            "第三部分 历史模拟法",
            chinese=True,
        )


if __name__ == "__main__":
    main()
